import pytest
from vulpcode.tools import get_tool
import vulpcode.tools.write_docx  # noqa: F401  (registers WriteDocx)

docx = pytest.importorskip("docx")


@pytest.mark.asyncio
async def test_write_docx_from_blocks(tmp_path):
    cls = get_tool("WriteDocx")
    target = tmp_path / "doc.docx"
    res = await cls().run(cls.Input(
        file_path=str(target),
        title="My Report",
        blocks=[
            {"kind": "heading", "text": "Introduction", "level": 1},
            {"kind": "paragraph", "text": "Hello world."},
            {"kind": "bullets", "text": ["Item A", "Item B"]},
            {"kind": "code", "text": "print('hi')", "language": "python"},
        ],
    ))
    assert res.is_error is False
    assert target.exists()
    # Re-open to confirm valid docx
    from docx import Document
    import io
    doc = Document(str(target))
    assert len(doc.paragraphs) > 0


@pytest.mark.asyncio
async def test_write_docx_from_markdown(tmp_path):
    cls = get_tool("WriteDocx")
    target = tmp_path / "md.docx"
    md = "# Heading\n\nSome paragraph.\n\n- bullet one\n- bullet two\n\n```python\nx = 1\n```\n"
    res = await cls().run(cls.Input(file_path=str(target), markdown=md))
    assert res.is_error is False
    assert target.exists()
    from docx import Document
    doc = Document(str(target))
    assert len(doc.paragraphs) > 0


@pytest.mark.asyncio
async def test_write_docx_missing_both_input(tmp_path):
    cls = get_tool("WriteDocx")
    with pytest.raises(Exception):
        cls.Input(file_path=str(tmp_path / "x.docx"))
