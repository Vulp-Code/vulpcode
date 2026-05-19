"""WriteDocx tool: build a .docx via python-docx with structured input."""
from __future__ import annotations

from typing import Literal

from pydantic import BaseModel, model_validator

from vulpcode.tools._validated_write import (
    ValidatedWriteTool,
    ValidationError,
)
from vulpcode.tools.base import tool


class _Block(BaseModel):
    kind: Literal["heading", "paragraph", "bullets", "numbered", "code"]
    text: str | list[str] = ""
    level: int = 1
    language: str = ""


def _markdown_to_blocks(md: str) -> list[_Block]:
    """Tiny MD->blocks converter (headings, paragraphs, bullets, fenced code)."""
    blocks: list[_Block] = []
    lines = md.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.startswith("```"):
            lang = line[3:].strip()
            code: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code.append(lines[i])
                i += 1
            blocks.append(_Block(kind="code", text="\n".join(code), language=lang))
            i += 1
            continue
        m = 0
        while m < len(line) and line[m] == "#":
            m += 1
        if 0 < m <= 6 and m < len(line) and line[m] == " ":
            blocks.append(_Block(kind="heading", text=line[m + 1:].strip(), level=m))
            i += 1
            continue
        if line.lstrip().startswith(("- ", "* ")):
            items: list[str] = []
            while i < len(lines) and lines[i].lstrip().startswith(("- ", "* ")):
                items.append(lines[i].lstrip()[2:])
                i += 1
            blocks.append(_Block(kind="bullets", text=items))
            continue
        if line.strip():
            para: list[str] = []
            while i < len(lines) and lines[i].strip():
                para.append(lines[i])
                i += 1
            blocks.append(_Block(kind="paragraph", text=" ".join(para)))
            continue
        i += 1
    return blocks


@tool(
    name="WriteDocx",
    description=(
        "Create or overwrite a Word document (.docx). Pass either `blocks` "
        "(structured: heading/paragraph/bullets/numbered/code) or `markdown` "
        "(converted internally). The file is validated by re-opening it with "
        "python-docx before commit."
    ),
    requires_confirm=True,
)
class WriteDocxTool(ValidatedWriteTool):
    binary = True

    class Input(BaseModel):
        file_path: str
        title: str = ""
        blocks: list[_Block] | None = None
        markdown: str | None = None

        @model_validator(mode="after")
        def _need_one(self):
            if self.blocks is None and self.markdown is None:
                raise ValueError("provide either `blocks` or `markdown`")
            return self

    def transform(self, args):
        try:
            from docx import Document
        except ImportError:
            raise ValidationError(
                "WriteDocx requires python-docx. Install with: pip install vulpcode[docs-tools]"
            )
        import io

        blocks = args.blocks if args.blocks is not None else _markdown_to_blocks(args.markdown or "")
        doc = Document()
        if args.title:
            doc.add_heading(args.title, level=0)
        for b in blocks:
            if b.kind == "heading":
                doc.add_heading(
                    b.text if isinstance(b.text, str) else " ".join(b.text),
                    level=b.level,
                )
            elif b.kind == "paragraph":
                doc.add_paragraph(b.text if isinstance(b.text, str) else " ".join(b.text))
            elif b.kind == "bullets":
                items = b.text if isinstance(b.text, list) else [b.text]
                for it in items:
                    doc.add_paragraph(it, style="List Bullet")
            elif b.kind == "numbered":
                items = b.text if isinstance(b.text, list) else [b.text]
                for it in items:
                    doc.add_paragraph(it, style="List Number")
            elif b.kind == "code":
                code = b.text if isinstance(b.text, str) else "\n".join(b.text)
                p = doc.add_paragraph()
                run = p.add_run(code)
                run.font.name = "Courier New"
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def validate(self, content, args):
        try:
            from docx import Document
            import io
            Document(io.BytesIO(content))
        except ImportError:
            raise ValidationError(
                "WriteDocx requires python-docx. Install with: pip install vulpcode[docs-tools]"
            )
        except Exception as e:
            raise ValidationError(f"Failed to re-open generated .docx: {e}")
